from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from docx import Document
import io
import os

from langchain_anthropic import ChatAnthropic
from langchain_core.prompts import ChatPromptTemplate
from dotenv import load_dotenv

load_dotenv()

app = Flask(__name__)
CORS(app)

# Megosztott adat a Streamlit és az editor között
shared_data = {
    "szoveg": "",
    "pontszam": 0
}

llm = ChatAnthropic(
    model="claude-haiku-4-5",
    api_key=os.getenv("ANTHROPIC_API_KEY")
)

@app.route("/get_text", methods=["GET"])
def get_text():
    return jsonify(shared_data)

@app.route("/set_text", methods=["POST"])
def set_text():
    data = request.json
    shared_data["szoveg"] = data.get("szoveg", "")
    shared_data["pontszam"] = data.get("pontszam", 0)
    return jsonify({"ok": True})

def rewrite_text(eredeti, utasitas, kontextus=""):
    """Szövegrészlet átírása AI segítségével. Közvetlenül hívható a PyQt6 appból is."""
    prompt = ChatPromptTemplate.from_messages([
        ("system", """Te egy profi magyar pályázatíró vagy.
        Csak a megadott részt írd át az utasítás alapján!
        Tartsd meg a pályázati stílust!
        Csak a kész szöveget add vissza, semmi mást!
        """),
        ("human", """
        Eredeti szövegrészlet:
        {eredeti}

        Utasítás: {utasitas}

        Kontextus: {kontextus}
        """)
    ])

    chain = prompt | llm
    response = chain.invoke({
        "eredeti": eredeti,
        "utasitas": utasitas,
        "kontextus": kontextus
    })
    return response.content.strip()


@app.route("/rewrite", methods=["POST"])
def rewrite():
    data = request.json
    uj_szoveg = rewrite_text(
        data.get("eredeti", ""),
        data.get("utasitas", ""),
        data.get("kontextus", "")
    )
    return jsonify({"uj_szoveg": uj_szoveg})

@app.route("/export_word", methods=["POST"])
def export_word():
    data = request.json
    szoveg = data.get("szoveg", "")

    doc = Document()
    doc.add_heading("Pályázati szöveg", 0)
    for bekezdes in szoveg.split("\n\n"):
        if bekezdes.strip():
            doc.add_paragraph(bekezdes.strip())

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return send_file(
        buffer,
        as_attachment=True,
        download_name="palyazat.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

if __name__ == "__main__":
    app.run(port=8502)

@app.route("/editor.html")
def editor():
    return send_file("editor.html")