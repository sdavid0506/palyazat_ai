import streamlit as st
from orchestrator import run
from file_reader import read_file
from rag_modul import add_document
from stylist_agent import analyze_and_save
from docx import Document
import tempfile
import os

st.set_page_config(
    page_title="AI Pályázatíró",
    page_icon="📝",
    layout="wide"
)

# Stílus
st.markdown("""
<style>
    .main-header {
        font-size: 2rem;
        font-weight: bold;
        color: #1f77b4;
        margin-bottom: 0.5rem;
    }
    .section-header {
        font-size: 1.2rem;
        font-weight: bold;
        margin-top: 1rem;
    }
</style>
""", unsafe_allow_html=True)

st.markdown('<div class="main-header">📝 AI Pályázatíró Rendszer</div>', unsafe_allow_html=True)
st.markdown("*Multi-agent rendszer professzionális pályázatok írásához*")
st.markdown("---")

# Oldalsáv
with st.sidebar:
    st.header("⚙️ Beállítások")
    max_rounds = st.slider(
        "Maximális javítási körök",
        min_value=1,
        max_value=5,
        value=2,
        help="Hányszor javítsa újra az AI a szöveget?"
    )
    st.markdown("---")
    st.markdown("### 🔄 Hogyan működik?")
    st.markdown("1. 📂 Töltsd fel a régi pályázatokat")
    st.markdown("2. 📋 Töltsd fel a kiírást")
    st.markdown("3. ✍️ Add meg a cégadatokat")
    st.markdown("4. 🚀 Az AI megírja a szöveget")
    st.markdown("5. ⬇️ Töltsd le Word formátumban")
    st.markdown("---")
    st.markdown("### 📊 Rendszer állapot")
    st.success("✅ AI motor: aktív")
    st.success("✅ Adatbázis: aktív")

# Három oszlop a feltöltéshez
col1, col2, col3 = st.columns(3)

with col1:
    st.markdown("### 📂 Régi pályázatok")
    st.caption("Ezekből tanulja a stílust az AI")
    regi_fajlok = st.file_uploader(
        "Húzd ide a régi pályázatokat",
        type=["pdf", "docx"],
        accept_multiple_files=True,
        key="regi"
    )

    if regi_fajlok:
        if st.button("📚 Betöltés AI memóriába", use_container_width=True):
            with st.spinner("Feldolgozás..."):
                for fajl in regi_fajlok:
                    with tempfile.NamedTemporaryFile(
                        delete=False,
                        suffix=os.path.splitext(fajl.name)[1]
                    ) as tmp:
                        tmp.write(fajl.read())
                        tmp_path = tmp.name

                    szoveg = read_file(tmp_path)
                    if szoveg:
                        analyze_and_save(szoveg, fajl.name)
                        st.success(f"✅ {fajl.name}")
                    os.unlink(tmp_path)

with col2:
    st.markdown("### 📋 Pályázati kiírás")
    st.caption("Az új kiírás amit be kell adni")
    kiiras_fajl = st.file_uploader(
        "Húzd ide a kiírást",
        type=["pdf", "docx"],
        key="kiiras"
    )

    tender_text = ""
    if kiiras_fajl:
        with tempfile.NamedTemporaryFile(
            delete=False,
            suffix=os.path.splitext(kiiras_fajl.name)[1]
        ) as tmp:
            tmp.write(kiiras_fajl.read())
            tmp_path = tmp.name

        tender_text = read_file(tmp_path)
        os.unlink(tmp_path)

        if tender_text:
            st.success(f"✅ {kiiras_fajl.name} beolvasva!")

with col3:
    st.markdown("### 🏢 Cégadatlap")
    st.caption("A cég és projekt adatai")
    ceg_fajl = st.file_uploader(
        "Húzd ide a cégadatlapot",
        type=["pdf", "docx"],
        key="ceg"
    )

    ceg_text = ""
    if ceg_fajl:
        with tempfile.NamedTemporaryFile(
            delete=False,
            suffix=os.path.splitext(ceg_fajl.name)[1]
        ) as tmp:
            tmp.write(ceg_fajl.read())
            tmp_path = tmp.name

        ceg_text = read_file(tmp_path)
        os.unlink(tmp_path)

        if ceg_text:
            st.success(f"✅ {ceg_fajl.name} beolvasva!")

st.markdown("---")

# Adatok megadása
st.markdown("### ✍️ Pályázat adatok")
col_bal, col_jobb = st.columns(2)

with col_bal:
    feladat = st.text_area(
        "Mit írjon meg az AI?",
        placeholder="Pl: Írj bevezető fejezetet egy digitális fejlesztési pályázathoz",
        height=120
    )

with col_jobb:
    adatok = st.text_area(
        "Cégadatok és projekt adatok",
        value=ceg_text if ceg_text else "",
        placeholder="""Pl:
Cég neve: TechMagyar Kft.
Projekt költsége: 5 000 000 Ft
Támogatás: 85%
Kezdés: 2025.01.01""",
        height=120
    )

st.markdown("---")

# Generálás gomb
if st.button("🚀 Pályázat generálása!", type="primary", use_container_width=True):
    if not feladat:
        st.error("❌ Kérlek add meg mit írjon az AI!")
    elif not adatok:
        st.error("❌ Kérlek add meg a cégadatokat!")
    else:
        # Progress bar
        progress = st.progress(0, text="Inicializálás...")

        with st.spinner("⏳ Az AI dolgozik..."):
            progress.progress(20, text="📋 Kiírás elemzése...")
            
            progress.progress(40, text="🎨 Stílus elemzése...")
            
            progress.progress(60, text="✍️ Szöveg generálása...")
            eredmeny, pontszam = run(
                task=feladat,
                data=adatok,
                tender_text=tender_text if tender_text else None,
                style_text=adatok,
                max_rounds=max_rounds
            )

            progress.progress(100, text="✅ Kész!")

        if eredmeny:
            st.success("✅ A pályázat elkészült!")
            st.markdown("### 📄 Generált szöveg:")
            st.text_area("Eredmény", eredmeny, height=400)

            # Word export
            doc = Document()
            doc.add_heading("Pályázati szöveg", 0)
            for bekezdes in eredmeny.split("\n\n"):
                if bekezdes.strip():
                    doc.add_paragraph(bekezdes.strip())

            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
                doc.save(tmp.name)
                with open(tmp.name, "rb") as f:
                    st.download_button(
                        label="⬇️ Letöltés Word formátumban",
                        data=f,
                        file_name="palyazat.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
                os.unlink(tmp.name)
        else:
            st.error("❌ Nem sikerült szöveget generálni, próbáld újra!")