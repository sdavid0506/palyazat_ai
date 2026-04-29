"""
text_beautifier.py
------------------
Markdown-szerű jelölések (##, **, * stb.) átalakítása
QTextEdit-kompatibilis HTML-re, valamint Qt HTML → .docx export.

Használat:
    from text_beautifier import beautify_to_html, build_docx_from_editor_html
    html = beautify_to_html(plain_text)
    editor.setHtml(html)

    build_docx_from_editor_html(editor.toHtml(), path)
"""

import re
import html as _html_mod
from html.parser import HTMLParser


# ── Markdown → HTML ───────────────────────────────────────────────────────────

def beautify_to_html(text: str) -> str:
    """
    Feldolgoz egy plain text / markdown-szerű szöveget és HTML-t ad vissza.

    Kezelt formák:
        # Cím         → <h1>
        ## Alcím      → <h2>
        ### Al-alcím  → <h3>
        **szöveg**    → félkövér
        *szöveg*      → dőlt
        _szöveg_      → dőlt
        - elem        → felsorolás (<ul><li>)
        üres sor      → bekezdésváltó
    """
    lines = text.splitlines()
    html_parts = []
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # H1 – csak egyszeres #, de NEM ## vagy ###
        m = re.match(r'^#{1}(?!#)\s+(.+)', stripped)
        if m:
            html_parts.append(
                f'<h1 style="font-size:18pt; font-weight:bold; margin:8px 0 4px 0;">'
                f'{_inline(m.group(1))}</h1>'
            )
            i += 1
            continue

        # H2
        m = re.match(r'^#{2}(?!#)\s+(.+)', stripped)
        if m:
            html_parts.append(
                f'<h2 style="font-size:14pt; font-weight:bold; margin:6px 0 3px 0;">'
                f'{_inline(m.group(1))}</h2>'
            )
            i += 1
            continue

        # H3
        m = re.match(r'^#{3}\s+(.+)', stripped)
        if m:
            html_parts.append(
                f'<h3 style="font-size:12pt; font-weight:bold; margin:4px 0 2px 0;">'
                f'{_inline(m.group(1))}</h3>'
            )
            i += 1
            continue

        # Vízszintes elválasztó: --- vagy *** vagy ___ → üres sor
        if re.match(r'^[-*_]{3,}$', stripped):
            html_parts.append('<p style="margin:4px 0;">&nbsp;</p>')
            i += 1
            continue

        # Táblázat: markdown pipe-szintaxis | Fejléc | ...
        if '|' in stripped and i + 1 < len(lines) and re.match(r'^\|[\s\-:|]+\|', lines[i + 1].strip()):
            header_cells = [_inline(c.strip()) for c in stripped.strip('|').split('|')]
            i += 2  # fejléc + elválasztó sor kihagyása
            data_rows = []
            while i < len(lines) and '|' in lines[i].strip():
                row_cells = [_inline(c.strip()) for c in lines[i].strip().strip('|').split('|')]
                data_rows.append(row_cells)
                i += 1
            th = "background:#f1f5f9; border:1px solid #cbd5e1; padding:6px 10px; font-weight:bold; text-align:left;"
            td = "border:1px solid #cbd5e1; padding:6px 10px;"
            thead = '<tr>' + ''.join(f'<th style="{th}">{c}</th>' for c in header_cells) + '</tr>'
            tbody = ''.join(
                '<tr>' + ''.join(f'<td style="{td}">{c}</td>' for c in row) + '</tr>'
                for row in data_rows
            )
            html_parts.append(
                '<table border="1" cellspacing="0" '
                'style="border-collapse:collapse; width:100%; margin:8px 0;">'
                f'<thead>{thead}</thead><tbody>{tbody}</tbody></table>'
            )
            continue

        # Lista: egymás utáni "- " vagy "* " sorokat összegyűjtjük
        if re.match(r'^[-*]\s+', stripped):
            items = []
            while i < len(lines) and re.match(r'^[-*]\s+', lines[i].strip()):
                item_text = re.match(r'^[-*]\s+(.+)', lines[i].strip()).group(1)
                items.append(f'<li style="margin:2px 0;">{_inline(item_text)}</li>')
                i += 1
            html_parts.append(
                '<ul style="margin:4px 0; padding-left:20px;">' + ''.join(items) + '</ul>'
            )
            continue

        # Üres sor → térköz
        if not stripped:
            html_parts.append('<p style="margin:4px 0;">&nbsp;</p>')
            i += 1
            continue

        # Normál bekezdés
        html_parts.append(
            f'<p style="margin:3px 0; line-height:1.6;">{_inline(stripped)}</p>'
        )
        i += 1

    body = '\n'.join(html_parts)
    return (
        '<html><body style="font-family: Segoe UI, sans-serif; font-size: 11pt; color: #1e293b;">'
        f'{body}'
        '</body></html>'
    )


def _inline(text: str) -> str:
    """Inline markdown → HTML (félkövér, dőlt)."""
    text = _html_mod.escape(text)
    text = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', text)
    text = re.sub(r'(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)', r'<i>\1</i>', text)
    text = re.sub(r'_(.+?)_', r'<i>\1</i>', text)
    return text


# ── Qt HTML → .docx ──────────────────────────────────────────────────────────

class _QtHtmlToDocx(HTMLParser):
    """
    Értelmezi a Qt QTextEdit.toHtml() kimenetét és python-docx bekezdéseket épít.
    Kezeli a <h1>–<h3>, <p>, <b>, <i>, <ul>, <li>, <span style="..."> elemeket.
    """

    _SKIP = {'html', 'body', 'meta', 'qt', 'br', 'thead', 'tbody', 'tfoot'}
    _IGNORE_CONTENT = {'head', 'style', 'script'}

    def __init__(self, doc):
        super().__init__()
        self.doc = doc
        self._para = None
        self._bold_stack = [False]
        self._italic_stack = [False]
        self._underline_stack = [False]
        self._heading = 0       # 0 = nem cím
        self._in_list = False
        self._in_li = False
        self._ignore_depth = 0  # >0: head/style belsejében vagyunk
        self._in_table = False
        self._table_rows = []         # [[(szöveg, is_header), ...], ...]
        self._table_current_row = []
        self._table_cell_buf = []
        self._in_td = False
        self._table_cell_is_header = False

    # ── stílus kinyerése ──────────────────────────────────────────────────

    @staticmethod
    def _parse_style(attrs):
        style = dict(attrs).get('style', '')
        bold = (
            'font-weight:700' in style
            or 'font-weight: 700' in style
            or 'font-weight:bold' in style
        )
        italic = 'font-style:italic' in style or 'font-style: italic' in style
        underline = 'text-decoration: underline' in style or 'text-decoration:underline' in style
        font_size = None
        m = re.search(r'font-size\s*:\s*(\d+)pt', style)
        if m:
            font_size = int(m.group(1))
        return bold, italic, underline, font_size

    # ── tag kezelők ──────────────────────────────────────────────────────

    def handle_starttag(self, tag, attrs):
        if tag in self._IGNORE_CONTENT:
            self._ignore_depth += 1
            return
        if tag in self._SKIP:
            return
        if self._ignore_depth > 0:
            return

        bold, italic, underline, font_size = self._parse_style(attrs)

        if tag in ('h1', 'h2', 'h3'):
            level = int(tag[1])
            self._heading = level
            self._para = self.doc.add_heading('', level=level)
            self._bold_stack.append(True)
            self._italic_stack.append(False)
            self._underline_stack.append(False)

        elif tag == 'p':
            if font_size and font_size >= 18:
                self._heading = 1
                self._para = self.doc.add_heading('', level=1)
            elif font_size and font_size >= 14:
                self._heading = 2
                self._para = self.doc.add_heading('', level=2)
            elif font_size and font_size >= 12:
                self._heading = 3
                self._para = self.doc.add_heading('', level=3)
            else:
                self._heading = 0
                self._para = self.doc.add_paragraph()
            self._bold_stack.append(bold)
            self._italic_stack.append(italic)
            self._underline_stack.append(underline)

        elif tag in ('b', 'strong'):
            self._bold_stack.append(True)
            self._italic_stack.append(self._italic_stack[-1])
            self._underline_stack.append(self._underline_stack[-1])

        elif tag in ('i', 'em'):
            self._bold_stack.append(self._bold_stack[-1])
            self._italic_stack.append(True)
            self._underline_stack.append(self._underline_stack[-1])

        elif tag == 'u':
            self._bold_stack.append(self._bold_stack[-1])
            self._italic_stack.append(self._italic_stack[-1])
            self._underline_stack.append(True)

        elif tag == 'span':
            self._bold_stack.append(bold or self._bold_stack[-1])
            self._italic_stack.append(italic or self._italic_stack[-1])
            self._underline_stack.append(underline or self._underline_stack[-1])

        elif tag == 'ul':
            self._in_list = True

        elif tag == 'li':
            self._in_li = True
            self._para = self.doc.add_paragraph(style='List Bullet')
            self._bold_stack.append(False)
            self._italic_stack.append(False)
            self._underline_stack.append(False)

        elif tag == 'table':
            self._in_table = True
            self._table_rows = []
            self._table_current_row = []
            self._table_cell_buf = []
            self._in_td = False

        elif tag == 'tr':
            if self._in_table:
                self._table_current_row = []

        elif tag in ('td', 'th'):
            if self._in_table:
                self._in_td = True
                self._table_cell_buf = []
                self._table_cell_is_header = (tag == 'th')

    def handle_endtag(self, tag):
        if tag in self._IGNORE_CONTENT:
            self._ignore_depth = max(0, self._ignore_depth - 1)
            return
        if tag in self._SKIP:
            return
        if self._ignore_depth > 0:
            return
        if tag in ('h1', 'h2', 'h3', 'p', 'b', 'strong', 'i', 'em', 'u', 'span', 'li'):
            if len(self._bold_stack) > 1:
                self._bold_stack.pop()
                self._italic_stack.pop()
                self._underline_stack.pop()
        if tag in ('h1', 'h2', 'h3'):
            self._heading = 0
        if tag == 'li':
            self._in_li = False
        if tag == 'ul':
            self._in_list = False
        if tag in ('td', 'th'):
            if self._in_table and self._in_td:
                self._table_current_row.append(
                    (list(self._table_cell_buf), self._table_cell_is_header)
                )
                self._in_td = False
        if tag == 'tr':
            if self._in_table and self._table_current_row:
                self._table_rows.append(self._table_current_row)
                self._table_current_row = []
        if tag == 'table':
            self._in_table = False
            if self._table_rows:
                num_cols = max(len(r) for r in self._table_rows)
                if num_cols > 0:
                    table = self.doc.add_table(rows=0, cols=num_cols)
                    table.style = 'Table Grid'
                    for row_data in self._table_rows:
                        row = table.add_row()
                        for ci, (runs, is_header) in enumerate(row_data):
                            if ci < num_cols:
                                para = row.cells[ci].paragraphs[0]
                                for (text, bold, italic, underline) in runs:
                                    r = para.add_run(text)
                                    r.bold = True if (bold or is_header) else None
                                    r.italic = italic or None
                                    r.underline = underline or None

    def handle_data(self, data):
        if self._ignore_depth > 0:
            return
        if self._in_table and self._in_td:
            self._table_cell_buf.append((
                data,
                self._bold_stack[-1],
                self._italic_stack[-1],
                self._underline_stack[-1],
            ))
            return
        if not data.strip():
            return
        if self._para is None:
            self._para = self.doc.add_paragraph()

        run = self._para.add_run(data)
        run.bold = self._bold_stack[-1] or None
        run.italic = self._italic_stack[-1] or None
        run.underline = self._underline_stack[-1] or None


def build_docx_from_editor_html(html_content: str, path: str):
    """
    Felépít egy .docx fájlt a QTextEdit.toHtml() kimenetéből,
    megőrizve a félkövér, dőlt, aláhúzott és heading formázásokat.
    """
    from docx import Document
    doc = Document()
    parser = _QtHtmlToDocx(doc)
    parser.feed(html_content)
    doc.save(path)
